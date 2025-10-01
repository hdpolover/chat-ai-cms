"""File extraction service for processing various document formats."""

import io
import mimetypes
import os
from typing import Dict, Any, Optional
import structlog

# File processing libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

logger = structlog.get_logger()


class FileExtractionService:
    """Service for extracting text content from various file formats."""

    def __init__(self):
        self.supported_extensions = {
            # Text files
            '.txt': self._extract_text,
            '.md': self._extract_text,
            '.py': self._extract_text,
            '.js': self._extract_text,
            '.ts': self._extract_text,
            '.json': self._extract_text,
            '.xml': self._extract_text,
            '.html': self._extract_text,
            '.css': self._extract_text,
            '.sql': self._extract_text,
            '.log': self._extract_text,
            '.yaml': self._extract_text,
            '.yml': self._extract_text,
            '.toml': self._extract_text,
            '.ini': self._extract_text,
            '.cfg': self._extract_text,
            '.conf': self._extract_text,
            
            # CSV files
            '.csv': self._extract_csv,
            '.tsv': self._extract_csv,
            
            # Excel files
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
            
            # PDF files
            '.pdf': self._extract_pdf,
            
            # Word documents
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,  # Will try to handle as docx
        }

    def extract_text_from_file(self, file_content: bytes, filename: str, mime_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text content from uploaded file.
        
        Returns:
            dict: {
                'content': str,
                'metadata': dict,
                'error': str (optional)
            }
        """
        try:
            # Get file extension
            _, ext = os.path.splitext(filename.lower())
            
            # Detect mime type if not provided
            if not mime_type:
                mime_type, _ = mimetypes.guess_type(filename)
            
            logger.info(
                "Processing file",
                filename=filename,
                extension=ext,
                mime_type=mime_type,
                file_size=len(file_content)
            )
            
            # Check if extension is supported
            if ext not in self.supported_extensions:
                return {
                    'content': '',
                    'metadata': {'original_filename': filename},
                    'error': f"Unsupported file type: {ext}. Supported types: {', '.join(sorted(self.supported_extensions.keys()))}"
                }
            
            # Extract content using appropriate method
            extraction_method = self.supported_extensions[ext]
            result = extraction_method(file_content, filename, mime_type)
            
            # Add common metadata
            result['metadata'].update({
                'original_filename': filename,
                'file_extension': ext,
                'mime_type': mime_type,
                'file_size_bytes': len(file_content),
                'extraction_method': extraction_method.__name__
            })
            
            return result
            
        except Exception as e:
            logger.error(
                "File extraction failed",
                filename=filename,
                error=str(e)
            )
            return {
                'content': '',
                'metadata': {'original_filename': filename},
                'error': f"Failed to process file {filename}: {str(e)}"
            }

    def _extract_text(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Extract content from plain text files."""
        try:
            # Try UTF-8 first
            content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                content = file_content.decode('latin-1')
            except UnicodeDecodeError:
                try:
                    # Last resort: utf-8 with error handling
                    content = file_content.decode('utf-8', errors='replace')
                except Exception:
                    raise Exception("Cannot decode text file - unsupported encoding")
        
        return {
            'content': content,
            'metadata': {
                'encoding': 'utf-8',
                'line_count': len(content.splitlines()),
                'character_count': len(content)
            }
        }

    def _extract_csv(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Extract content from CSV/TSV files."""
        if pd is None:
            raise Exception("pandas library not available for CSV processing")
        
        try:
            # Determine separator
            separator = '\t' if filename.lower().endswith('.tsv') else ','
            
            # Read CSV with pandas
            df = pd.read_csv(
                io.BytesIO(file_content),
                sep=separator,
                encoding='utf-8',
                on_bad_lines='skip'  # Skip problematic lines
            )
            
            # Convert to readable text format
            content_parts = []
            
            # Add column headers
            content_parts.append("Column Headers:")
            content_parts.append(", ".join(df.columns.tolist()))
            content_parts.append("")
            
            # Add sample data (first few rows)
            sample_rows = min(10, len(df))
            if sample_rows > 0:
                content_parts.append(f"Sample Data (showing first {sample_rows} rows):")
                for idx, row in df.head(sample_rows).iterrows():
                    row_text = " | ".join([f"{col}: {str(val)}" for col, val in row.items()])
                    content_parts.append(f"Row {idx + 1}: {row_text}")
            
            # Add summary statistics for numeric columns
            numeric_columns = df.select_dtypes(include=['number']).columns
            if len(numeric_columns) > 0:
                content_parts.append("")
                content_parts.append("Numeric Column Summary:")
                for col in numeric_columns:
                    stats = df[col].describe()
                    content_parts.append(f"{col}: mean={stats['mean']:.2f}, min={stats['min']}, max={stats['max']}")
            
            content = "\n".join(content_parts)
            
            return {
                'content': content,
                'metadata': {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist(),
                    'data_types': df.dtypes.astype(str).to_dict(),
                    'separator': separator,
                    'numeric_columns': numeric_columns.tolist()
                }
            }
            
        except Exception as e:
            # Fallback to plain text extraction
            logger.warning(f"CSV parsing failed, falling back to text: {e}")
            return self._extract_text(file_content, filename, mime_type)

    def _extract_excel(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Extract content from Excel files."""
        if pd is None or openpyxl is None:
            raise Exception("pandas and openpyxl libraries not available for Excel processing")
        
        try:
            # Read Excel file with pandas
            excel_data = pd.read_excel(
                io.BytesIO(file_content),
                sheet_name=None,  # Read all sheets
                engine='openpyxl'
            )
            
            content_parts = []
            sheet_info = {}
            
            for sheet_name, df in excel_data.items():
                content_parts.append(f"=== Sheet: {sheet_name} ===")
                
                # Add column headers
                content_parts.append("Column Headers:")
                content_parts.append(", ".join(df.columns.tolist()))
                content_parts.append("")
                
                # Add sample data (first few rows)
                sample_rows = min(5, len(df))
                if sample_rows > 0:
                    content_parts.append(f"Sample Data (showing first {sample_rows} rows):")
                    for idx, row in df.head(sample_rows).iterrows():
                        row_text = " | ".join([f"{col}: {str(val)}" for col, val in row.items()])
                        content_parts.append(f"Row {idx + 1}: {row_text}")
                
                content_parts.append("")
                
                # Store sheet metadata
                sheet_info[sheet_name] = {
                    'row_count': len(df),
                    'column_count': len(df.columns),
                    'columns': df.columns.tolist()
                }
            
            content = "\n".join(content_parts)
            
            return {
                'content': content,
                'metadata': {
                    'sheet_count': len(excel_data),
                    'sheet_names': list(excel_data.keys()),
                    'sheets': sheet_info,
                    'total_rows': sum(len(df) for df in excel_data.values()),
                    'engine': 'openpyxl'
                }
            }
            
        except Exception as e:
            raise Exception(f"Failed to process Excel file: {str(e)}")

    def _extract_pdf(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Extract content from PDF files."""
        if PyPDF2 is None:
            raise Exception("PyPDF2 library not available for PDF processing")
        
        try:
            # Create PDF reader from bytes
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            content_parts = []
            page_count = len(pdf_reader.pages)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content_parts.append(f"=== Page {page_num + 1} ===")
                        content_parts.append(page_text)
                        content_parts.append("")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    content_parts.append(f"=== Page {page_num + 1} (extraction failed) ===")
                    content_parts.append("")
            
            content = "\n".join(content_parts)
            
            # Get PDF metadata
            pdf_metadata = {}
            if pdf_reader.metadata:
                try:
                    for key, value in pdf_reader.metadata.items():
                        if value:
                            pdf_metadata[key.replace('/', '')] = str(value)
                except Exception:
                    pass
            
            return {
                'content': content,
                'metadata': {
                    'page_count': page_count,
                    'pdf_metadata': pdf_metadata,
                    'extractor': 'PyPDF2',
                    'character_count': len(content)
                }
            }
            
        except Exception as e:
            raise Exception(f"Failed to process PDF file: {str(e)}")

    def _extract_docx(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Extract content from Word documents."""
        if DocxDocument is None:
            raise Exception("python-docx library not available for Word document processing")
        
        try:
            # Create document from bytes
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
            
            content_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                content_parts.append(f"\n=== Table {table_count} ===")
                
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_cells.append(cell_text)
                    
                    if any(row_cells):  # Only add non-empty rows
                        content_parts.append(" | ".join(row_cells))
                
                content_parts.append("")
            
            content = "\n".join(content_parts)
            
            # Extract document properties
            doc_properties = {}
            try:
                core_props = doc.core_properties
                if core_props.author:
                    doc_properties['author'] = core_props.author
                if core_props.title:
                    doc_properties['title'] = core_props.title
                if core_props.subject:
                    doc_properties['subject'] = core_props.subject
                if core_props.created:
                    doc_properties['created'] = str(core_props.created)
                if core_props.modified:
                    doc_properties['modified'] = str(core_props.modified)
            except Exception:
                pass
            
            return {
                'content': content,
                'metadata': {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': table_count,
                    'document_properties': doc_properties,
                    'character_count': len(content),
                    'extractor': 'python-docx'
                }
            }
            
        except Exception as e:
            # For .doc files, the error might be expected
            if filename.lower().endswith('.doc'):
                raise Exception("Legacy .doc files are not supported. Please convert to .docx format.")
            else:
                raise Exception(f"Failed to process Word document: {str(e)}")

    def get_supported_formats(self) -> Dict[str, str]:
        """Get a dictionary of supported file formats and their descriptions."""
        return {
            # Text files
            '.txt': 'Plain text files',
            '.md': 'Markdown files',
            '.py': 'Python source code',
            '.js': 'JavaScript files',
            '.ts': 'TypeScript files',
            '.json': 'JSON data files',
            '.xml': 'XML files',
            '.html': 'HTML files',
            '.css': 'CSS stylesheets',
            '.sql': 'SQL scripts',
            '.log': 'Log files',
            '.yaml/.yml': 'YAML configuration files',
            '.toml': 'TOML configuration files',
            '.ini/.cfg/.conf': 'Configuration files',
            
            # Data files
            '.csv/.tsv': 'Comma/Tab separated values',
            '.xlsx/.xls': 'Microsoft Excel spreadsheets',
            
            # Documents
            '.pdf': 'PDF documents (text extraction)',
            '.docx': 'Microsoft Word documents (.doc not supported)',
        }

    def validate_file_size(self, file_size: int, max_size_mb: int = 50) -> bool:
        """Validate if file size is within acceptable limits."""
        max_size_bytes = max_size_mb * 1024 * 1024
        return file_size <= max_size_bytes

    def get_file_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Get basic file information without extracting content."""
        file_size = len(file_content)
        _, ext = os.path.splitext(filename.lower())
        mime_type, _ = mimetypes.guess_type(filename)
        
        extraction_method = 'Unknown'
        if ext in self.supported_extensions:
            method = self.supported_extensions[ext]
            extraction_method = method.__name__ if hasattr(method, '__name__') else 'Unknown'
        
        return {
            'filename': filename,
            'extension': ext,
            'mime_type': mime_type,
            'size_bytes': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'is_supported': ext in self.supported_extensions,
            'extraction_method': extraction_method
        }


# Global instance
file_extraction_service = FileExtractionService()